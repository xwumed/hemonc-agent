#!/usr/bin/env python3
from pathlib import Path
import chromadb
from docx import Document as DocxDocument

from config_manager import (
    embed_client,
    EMBED_MODEL,
    PATHO_DIR,
    PATHO_DB_STORAGE,
)

# ChromaDB setup
STORAGE_PATH = str(PATHO_DB_STORAGE)
chroma_client = chromadb.PersistentClient(path=STORAGE_PATH)
collection = chroma_client.get_or_create_collection(name="patho_collection")

# ─────────── Helpers ───────────
def docx_to_text(docx_path: str) -> str:
    """
    Extract text from a .docx file, including paragraphs, headings, and tables.
    """
    try:
        doc = DocxDocument(docx_path)
        parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text.strip())

        return "\n\n".join(parts)
    except Exception as e:
        print(f"[error] Failed to parse {docx_path}: {e}")
        return ""

def split_text(text: str, max_tokens: int = 8000) -> list[str]:
    """
    Split text into chunks with approximately max_tokens tokens.
    Splits on paragraphs (double newlines) to preserve context.
    """
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = []
    current_length = 0

    # Rough token estimation: 1 token ≈ 4 characters
    for para in paragraphs:
        para_length = len(para) // 4 + 1  # Approximate tokens
        if current_length + para_length > max_tokens:
            if current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_length = 0
        current_chunk.append(para)
        current_length += para_length

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks

# ─────────── Main ───────────
def main():
    docx_dir = Path(PATHO_DIR)

    documents = []
    metadatas = []
    ids = []

    for docx_path in sorted(docx_dir.glob("*.docx")):
        stem = docx_path.stem
        print(f"[info] Processing {docx_path.name}...")

        # 1) Extract text from .docx
        text = docx_to_text(docx_path)
        if not text:
            print(f"[warn] No text extracted from {docx_path.name} → skipping")
            continue

        # 2) Split into chunks if necessary
        chunks = split_text(text, max_tokens=8000)
        for i, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            documents.append(chunk)
            metadatas.append({"docx_name": stem, "chunk_index": i})
            ids.append(f"{stem}_chunk_{i}")

    if not documents:
        print("No documents found—check your folder.")
        return

    # ─────────── Embed & Persist ───────────
    print(f"\n🔑 Embedding {len(documents)} document chunks...")
    embeddings = []
    for text in documents:
        try:
            resp = embed_client.embeddings.create(
                model=EMBED_MODEL,
                input=[text],
                encoding_format="float"
            )
            embeddings.append(resp.data[0].embedding)
        except Exception as e:
            print(f"[error] Failed to embed document chunk: {e}")
            embeddings.append(None)

    # Filter out any failed embeddings
    valid_indices = [i for i, emb in enumerate(embeddings) if emb is not None]
    valid_documents = [documents[i] for i in valid_indices]
    valid_embeddings = [embeddings[i] for i in valid_indices]
    valid_metadatas = [metadatas[i] for i in valid_indices]
    valid_ids = [ids[i] for i in valid_indices]

    if valid_documents:
        collection.upsert(
            documents=valid_documents,
            embeddings=valid_embeddings,
            metadatas=valid_metadatas,
            ids=valid_ids
        )
        print(f"✅ Indexed {len(valid_documents)} document chunks into Chroma at {STORAGE_PATH!r}")
    else:
        print("No valid documents to index.")

if __name__ == "__main__":
    main()